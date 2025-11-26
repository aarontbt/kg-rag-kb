from docx import Document
from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class DOCXProcessor:
    """Process Word documents to extract text and metadata"""
    
    def __init__(self):
        pass
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a Word document and return extracted text with metadata
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            List of dictionaries containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        try:
            doc = Document(file_path)
            
            # Extract full document text
            full_text = self._extract_document_text(doc)
            
            # Extract text by sections (paragraphs)
            paragraphs = self._extract_paragraphs(doc)
            
            chunks = []
            
            # Add full document as one chunk
            if full_text:
                chunks.append({
                    'text': full_text.strip(),
                    'metadata': {
                        'source': str(file_path),
                        'file_type': 'docx',
                        'content_type': 'full_document',
                        'paragraph_count': len(paragraphs)
                    }
                })
            
            # Add individual paragraphs as separate chunks
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    chunks.append({
                        'text': paragraph.strip(),
                        'metadata': {
                            'source': str(file_path),
                            'file_type': 'docx',
                            'content_type': 'paragraph',
                            'paragraph_number': i + 1
                        }
                    })
            
            # Extract table content
            tables = self._extract_tables(doc)
            for i, table_text in enumerate(tables):
                if table_text.strip():
                    chunks.append({
                        'text': table_text.strip(),
                        'metadata': {
                            'source': str(file_path),
                            'file_type': 'docx',
                            'content_type': 'table',
                            'table_number': i + 1
                        }
                    })
            
            logger.info(f"Successfully processed Word document: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return []
    
    def _extract_document_text(self, doc: Document) -> str:
        """
        Extract the full text content from the document
        
        Args:
            doc: Word document object
            
        Returns:
            Full text content of the document
        """
        text_parts = []
        
        # Extract all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract table content
        for table in doc.tables:
            table_text = self._extract_single_table_text(table)
            if table_text.strip():
                text_parts.append(table_text)
        
        return '\n'.join(text_parts)
    
    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """
        Extract text from all paragraphs
        
        Args:
            doc: Word document object
            
        Returns:
            List of paragraph texts
        """
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[str]:
        """
        Extract text from all tables
        
        Args:
            doc: Word document object
            
        Returns:
            List of table texts
        """
        tables = []
        
        for table in doc.tables:
            table_text = self._extract_single_table_text(table)
            if table_text.strip():
                tables.append(table_text)
        
        return tables
    
    def _extract_single_table_text(self, table) -> str:
        """
        Extract text from a single table
        
        Args:
            table: Word table object
            
        Returns:
            Table text in CSV-like format
        """
        rows_text = []
        
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells_text.append(cell_text)
            
            if any(cells_text):  # Skip empty rows
                rows_text.append(' | '.join(cells_text))
        
        return '\n'.join(rows_text)
    
    def is_supported(self, file_path: str) -> bool:
        """Check if the file is a supported Word document format"""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
